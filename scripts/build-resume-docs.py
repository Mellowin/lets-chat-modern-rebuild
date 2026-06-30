#!/usr/bin/env python3
"""Generate DOCX and print-ready HTML from a markdown resume."""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

ROOT = Path(__file__).resolve().parent.parent
FINAL_DIR = ROOT / "docs" / "career" / "final"


def parse_inline(text):
    """Split text into (text, is_bold, link_url) segments."""
    segments = []
    pos = 0
    # Pattern for links [text](url) and bold **text**
    pattern = re.compile(r"\[([^\]]+)\]\(([^)]+)\)|\*\*([^*]+)\*\*")
    for m in pattern.finditer(text):
        if m.start() > pos:
            segments.append((text[pos : m.start()], False, None))
        if m.group(1) is not None:
            segments.append((m.group(1), False, m.group(2)))
        else:
            segments.append((m.group(3), True, None))
        pos = m.end()
    if pos < len(text):
        segments.append((text[pos:], False, None))
    return segments


def add_run(paragraph, text, bold=False, link=None, font_size=None, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    if link:
        run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
        run.underline = True
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = color
    return run


def md_to_docx(md_path, docx_path):
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(3)
    style.paragraph_format.space_before = Pt(0)

    # Page margins for 2-page fit
    sections = doc.sections[0]
    sections.top_margin = Inches(0.5)
    sections.bottom_margin = Inches(0.5)
    sections.left_margin = Inches(0.6)
    sections.right_margin = Inches(0.6)

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.read().splitlines()

    in_list = False
    list_items = []

    def flush_list():
        nonlocal in_list, list_items
        for content in list_items:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            for text, bold, link in parse_inline(content):
                add_run(p, text, bold=bold, link=link)
        in_list = False
        list_items = []

    for raw_line in lines:
        line = raw_line.rstrip()
        if not line:
            flush_list()
            continue

        # Headings
        if line.startswith("# "):
            flush_list()
            p = doc.add_heading(line[2:], level=0)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            for run in p.runs:
                run.font.size = Pt(22)
                run.font.bold = True
            continue
        if line.startswith("## "):
            flush_list()
            p = doc.add_heading(line[3:], level=1)
            for run in p.runs:
                run.font.size = Pt(13)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            continue
        if line.startswith("### "):
            flush_list()
            p = doc.add_heading(line[4:], level=2)
            for run in p.runs:
                run.font.size = Pt(12)
                run.font.bold = True
            continue

        # Horizontal rule
        if line.strip() == "---":
            flush_list()
            continue

        # Bullet list item
        if line.startswith("- "):
            list_items.append(line[2:])
            in_list = True
            continue

        flush_list()

        # Role line right under name
        if line.startswith("**") and line.endswith("**") and len(line) < 80:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            add_run(p, line[2:-2], bold=True, font_size=12)
            continue

        # Regular paragraph with inline formatting
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        for text, bold, link in parse_inline(line):
            add_run(p, text, bold=bold, link=link)

    flush_list()
    doc.save(docx_path)


def md_to_html(md_path, html_path, title):
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.read().splitlines()

    body_parts = []
    in_list = False

    def flush_list():
        nonlocal in_list
        if in_list:
            body_parts.append("</ul>")
            in_list = False

    for raw_line in lines:
        line = raw_line.rstrip()
        if not line:
            flush_list()
            continue
        if line.strip() == "---":
            flush_list()
            continue

        if line.startswith("# "):
            flush_list()
            body_parts.append(f'<h1>{escape(line[2:])}</h1>')
        elif line.startswith("## "):
            flush_list()
            body_parts.append(f'<h2>{escape(line[3:])}</h2>')
        elif line.startswith("### "):
            flush_list()
            body_parts.append(f'<h3>{escape(line[4:])}</h3>')
        elif line.startswith("- "):
            if not in_list:
                body_parts.append("<ul>")
                in_list = True
            body_parts.append(f'<li>{inline_html(line[2:])}</li>')
        else:
            flush_list()
            body_parts.append(f'<p>{inline_html(line)}</p>')

    flush_list()

    html = f"""<!DOCTYPE html>
<html lang="{ 'uk' if '_UA' in str(md_path) else 'en' }">
<head>
<meta charset="UTF-8">
<title>{escape(title)}</title>
<style>
  @page {{ size: A4; margin: 14mm 16mm; }}
  body {{
    font-family: "Segoe UI", Arial, sans-serif;
    font-size: 11pt;
    line-height: 1.35;
    color: #111;
    max-width: 180mm;
    margin: 0 auto;
  }}
  h1 {{ font-size: 22pt; margin: 0 0 2pt 0; }}
  h2 {{ font-size: 12.5pt; margin: 10pt 0 3pt 0; border-bottom: 1pt solid #ccc; padding-bottom: 2pt; }}
  h3 {{ font-size: 11.5pt; margin: 8pt 0 2pt 0; }}
  p {{ margin: 0 0 4pt 0; }}
  ul {{ margin: 2pt 0 6pt 16pt; padding: 0; }}
  li {{ margin-bottom: 2pt; }}
  a {{ color: #000; text-decoration: none; }}
  .role {{ font-weight: bold; font-size: 12pt; margin-bottom: 6pt; }}
</style>
</head>
<body>
{''.join(body_parts)}
</body>
</html>"""

    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html)


def escape(text):
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def inline_html(text):
    # bold
    text = re.sub(r"\*\*([^*]+)\*\*", r"<strong>\1</strong>", text)
    # links
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r'<a href="\2">\1</a>', text)
    return text


def main():
    resumes = [
        (FINAL_DIR / "Valerii_Khoidas_CV_UA.md", "Валерій Хойдас — CV"),
        (FINAL_DIR / "Valerii_Khoidas_CV_EN.md", "Valerii Khoidas — CV"),
    ]

    for md_path, title in resumes:
        if not md_path.exists():
            print(f"Missing {md_path}", file=sys.stderr)
            sys.exit(1)
        base = md_path.stem
        docx_path = FINAL_DIR / f"{base}.docx"
        html_path = FINAL_DIR / f"{base}.html"
        md_to_docx(md_path, docx_path)
        md_to_html(md_path, html_path, title)
        print(f"Generated: {docx_path.name}, {html_path.name}")


if __name__ == "__main__":
    main()
